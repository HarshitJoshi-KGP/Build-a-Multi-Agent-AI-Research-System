import streamlit as st
import time
import io
import re
from datetime import datetime

# ── Page Config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ResearchMind",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Export Helpers ────────────────────────────────────────────────────────────
def generate_pdf(topic: str, report: str, critique: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )

    ORANGE = colors.HexColor("#ff8c32")
    DARK   = colors.HexColor("#0a0a0f")
    GRAY   = colors.HexColor("#888888")
    WHITE  = colors.white

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("rm_title", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=28, leading=34,
        textColor=colors.HexColor("#1a1a2e"), spaceAfter=6)
    eyebrow_style = ParagraphStyle("rm_eyebrow", parent=styles["Normal"],
        fontName="Helvetica", fontSize=8, leading=12,
        textColor=ORANGE, spaceBefore=0, spaceAfter=12,
        wordWrap="CJK", letterSpacing=2)
    h2_style = ParagraphStyle("rm_h2", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=14, leading=18,
        textColor=colors.HexColor("#1a1a2e"), spaceBefore=18, spaceAfter=8)
    h3_style = ParagraphStyle("rm_h3", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=11, leading=14,
        textColor=colors.HexColor("#333333"), spaceBefore=12, spaceAfter=6)
    body_style = ParagraphStyle("rm_body", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, leading=16,
        textColor=colors.HexColor("#333333"), spaceAfter=8)
    meta_style = ParagraphStyle("rm_meta", parent=styles["Normal"],
        fontName="Helvetica", fontSize=8, leading=12,
        textColor=GRAY, spaceAfter=4)
    critique_style = ParagraphStyle("rm_critique", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9.5, leading=15,
        textColor=colors.HexColor("#444444"), spaceAfter=6,
        leftIndent=12, borderPad=8,
        backColor=colors.HexColor("#fdf6ee"))

    def md_to_paragraph(text, style):
        """Basic markdown → ReportLab paragraph (bold, headers)."""
        # Bold: **text** → <b>text</b>
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        # Escape ampersands not in tags
        text = re.sub(r'&(?!amp;|lt;|gt;|#)', '&amp;', text)
        return Paragraph(text, style)

    def parse_markdown_sections(text):
        """Parse markdown into (level, text) tuples."""
        elements = []
        for line in text.split('\n'):
            line = line.rstrip()
            if line.startswith('### '):
                elements.append(('h3', line[4:]))
            elif line.startswith('## '):
                elements.append(('h2', line[3:]))
            elif line.startswith('# '):
                elements.append(('h1', line[2:]))
            elif line.startswith('- ') or line.startswith('* '):
                elements.append(('bullet', line[2:]))
            elif line.strip() == '':
                elements.append(('space', ''))
            else:
                elements.append(('body', line))
        return elements

    story = []

    # ── Header Block ──
    story.append(Paragraph("INTELLIGENCE SYNTHESIS REPORT", eyebrow_style))
    story.append(Paragraph("ResearchMind", title_style))
    story.append(HRFlowable(width="100%", thickness=3, color=ORANGE, spaceAfter=16))

    # Meta table
    ts = datetime.now().strftime("%d %b %Y, %H:%M UTC")
    meta_data = [
        ["RESEARCH OBJECTIVE", topic],
        ["GENERATED", ts],
        ["PIPELINE", "Search → Reader → Writer → Critic"],
        ["ENGINE", "ResearchMind v2.0 · Multi-Agent Architecture"],
    ]
    meta_tbl = Table(meta_data, colWidths=[4*cm, 12*cm])
    meta_tbl.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
        ("FONTSIZE", (0,0), (-1,-1), 8),
        ("TEXTCOLOR", (0,0), (0,-1), ORANGE),
        ("TEXTCOLOR", (1,0), (1,-1), colors.HexColor("#333333")),
        ("TOPPADDING", (0,0), (-1,-1), 3),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
    ]))
    story.append(meta_tbl)
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5,
                            color=colors.HexColor("#dddddd"), spaceAfter=20))

    # ── Report Body ──
    story.append(Paragraph("SYNTHESIS REPORT", eyebrow_style))
    for kind, text in parse_markdown_sections(report):
        if kind == 'h1':
            story.append(Paragraph(text, title_style))
        elif kind == 'h2':
            story.append(Paragraph(text, h2_style))
        elif kind == 'h3':
            story.append(Paragraph(text, h3_style))
        elif kind == 'bullet':
            story.append(Paragraph(f"• {text}", body_style))
        elif kind == 'space':
            story.append(Spacer(1, 6))
        else:
            if text.strip():
                story.append(md_to_paragraph(text, body_style))

    # ── Critique Section ──
    if critique:
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=colors.HexColor("#dddddd"), spaceAfter=16))
        story.append(Paragraph("PEER REVIEW & CRITIQUE", eyebrow_style))
        for kind, text in parse_markdown_sections(critique):
            if kind in ('h2', 'h3'):
                story.append(Paragraph(text, h3_style))
            elif kind == 'bullet':
                story.append(Paragraph(f"• {text}", critique_style))
            elif kind == 'space':
                story.append(Spacer(1, 4))
            else:
                if text.strip():
                    story.append(md_to_paragraph(text, critique_style))

    doc.build(story)
    return buffer.getvalue()


def generate_docx(topic: str, report: str, critique: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    ORANGE = RGBColor(0xFF, 0x8C, 0x32)
    DARK   = RGBColor(0x1a, 0x1a, 0x2e)
    GRAY   = RGBColor(0x88, 0x88, 0x88)

    def set_font(run, size, bold=False, color=None, italic=False):
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        run.font.name = "Calibri"

    def add_eyebrow(doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        set_font(run, 8, bold=True, color=ORANGE)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_rule(doc, color_hex="#ff8c32", thickness=18):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(thickness))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex.lstrip('#'))
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(12)
        return p

    def add_thin_rule(doc):
        return add_rule(doc, "#dddddd", 4)

    def parse_and_add(doc, text, base_size=10.5):
        for line in text.split('\n'):
            line = line.rstrip()
            if line.startswith('### '):
                p = doc.add_paragraph()
                r = p.add_run(line[4:])
                set_font(r, 12, bold=True, color=DARK)
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
            elif line.startswith('## '):
                p = doc.add_paragraph()
                r = p.add_run(line[3:])
                set_font(r, 15, bold=True, color=DARK)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
            elif line.startswith('# '):
                p = doc.add_paragraph()
                r = p.add_run(line[2:])
                set_font(r, 20, bold=True, color=DARK)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(8)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                parts = re.split(r'\*\*(.+?)\*\*', line[2:])
                for i, part in enumerate(parts):
                    r = p.add_run(part)
                    set_font(r, base_size, bold=(i % 2 == 1))
            elif line.strip() == '':
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
            else:
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.+?)\*\*', line)
                for i, part in enumerate(parts):
                    r = p.add_run(part)
                    set_font(r, base_size, bold=(i % 2 == 1))
                p.paragraph_format.space_after = Pt(6)

    # ── Header ──
    add_eyebrow(doc, "Intelligence Synthesis Report")
    p = doc.add_paragraph()
    r = p.add_run("ResearchMind")
    set_font(r, 32, bold=True, color=DARK)
    p.paragraph_format.space_after = Pt(6)
    add_rule(doc)

    # Meta block
    ts = datetime.now().strftime("%d %b %Y, %H:%M UTC")
    meta = [
        ("RESEARCH OBJECTIVE", topic),
        ("GENERATED", ts),
        ("PIPELINE", "Search → Reader → Writer → Critic"),
        ("ENGINE", "ResearchMind v2.0 · Multi-Agent Architecture"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}:  ")
        set_font(r1, 8.5, bold=True, color=ORANGE)
        r2 = p.add_run(value)
        set_font(r2, 8.5, color=RGBColor(0x33, 0x33, 0x33))
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    add_thin_rule(doc)

    # ── Report ──
    add_eyebrow(doc, "Synthesis Report")
    parse_and_add(doc, report)

    # ── Critique ──
    if critique:
        doc.add_paragraph()
        add_thin_rule(doc)
        add_eyebrow(doc, "Peer Review & Critique")
        parse_and_add(doc, critique, base_size=10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Session State ─────────────────────────────────────────────────────────────
for key, default in [("results", {}), ("stage", "idle"), ("topic_submitted", "")]:
    if key not in st.session_state:
        st.session_state[key] = default

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Mono:ital,wght@0,300;0,400;0,500;1,400&family=DM+Sans:ital,wght@0,300;0,400;0,500;1,300&display=swap');

:root {
    --orange: #ff8c32;
    --orange-dim: rgba(255,140,50,0.15);
    --orange-glow: rgba(255,140,50,0.25);
    --green: #4ade80;
    --green-dim: rgba(74,222,128,0.12);
    --bg: #07070d;
    --surface: rgba(255,255,255,0.035);
    --surface-hover: rgba(255,255,255,0.06);
    --border: rgba(255,255,255,0.08);
    --border-lit: rgba(255,255,255,0.14);
    --text: #e2ddd4;
    --muted: #666;
    --muted2: #444;
}

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
    color: var(--text);
}

.stApp {
    background: var(--bg);
    background-image:
        radial-gradient(ellipse 60% 50% at 10% 0%, rgba(255,140,50,0.07) 0%, transparent 60%),
        radial-gradient(ellipse 40% 60% at 90% 100%, rgba(255,60,20,0.04) 0%, transparent 60%);
    min-height: 100vh;
}

#MainMenu, footer, header { visibility: hidden; }

.block-container {
    padding: 3rem 4rem 4rem;
    max-width: 1400px;
    margin: 0 auto;
}

/* ── Top nav bar ── */
.topbar {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 3.5rem;
    padding-bottom: 1.5rem;
    border-bottom: 1px solid var(--border);
}
.topbar-logo {
    font-family: 'Syne', sans-serif;
    font-weight: 800;
    font-size: 1.1rem;
    letter-spacing: -0.02em;
    color: var(--text);
}
.topbar-logo span { color: var(--orange); }
.topbar-badge {
    font-family: 'DM Mono', monospace;
    font-size: 0.6rem;
    letter-spacing: 0.18em;
    color: var(--muted);
    text-transform: uppercase;
    border: 1px solid var(--border);
    padding: 0.3rem 0.7rem;
    border-radius: 20px;
}

/* ── Hero ── */
.hero-eyebrow {
    font-family: 'DM Mono', monospace;
    font-size: 0.65rem;
    letter-spacing: 0.3em;
    text-transform: uppercase;
    color: var(--orange);
    margin-bottom: 1.2rem;
}
.hero-title {
    font-family: 'Syne', sans-serif;
    font-size: clamp(3rem, 6vw, 5.5rem);
    font-weight: 800;
    line-height: 0.88;
    letter-spacing: -0.05em;
    color: var(--text);
    margin-bottom: 1.4rem;
}
.hero-title .accent { color: var(--orange); }
.hero-sub {
    font-size: 1rem;
    line-height: 1.65;
    color: var(--muted);
    max-width: 520px;
    margin-bottom: 2.8rem;
}

/* ── Input area ── */
.input-wrapper {
    background: var(--surface);
    border: 1px solid var(--border-lit);
    border-radius: 18px;
    padding: 1.8rem 2rem;
    margin-bottom: 1rem;
    transition: border-color 0.3s;
}
.input-wrapper:focus-within {
    border-color: var(--orange);
    box-shadow: 0 0 0 3px var(--orange-glow), 0 12px 40px rgba(0,0,0,0.3);
}
.input-label {
    font-family: 'DM Mono', monospace;
    font-size: 0.6rem;
    letter-spacing: 0.25em;
    color: var(--orange);
    text-transform: uppercase;
    margin-bottom: 0.7rem;
}

.stTextInput > div > div {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    padding: 0 !important;
}
.stTextInput input {
    background: transparent !important;
    border: none !important;
    outline: none !important;
    box-shadow: none !important;
    color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 1.05rem !important;
    padding: 0 !important;
    caret-color: var(--orange);
}
.stTextInput input::placeholder { color: var(--muted2) !important; }

/* ── Primary button ── */
.stButton > button {
    background: var(--orange) !important;
    border: none !important;
    color: #0a0505 !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
    letter-spacing: 0.04em !important;
    height: 3.2rem !important;
    border-radius: 12px !important;
    width: 100%;
    cursor: pointer;
    transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1) !important;
}
.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 25px var(--orange-glow) !important;
    filter: brightness(1.08) !important;
}
.stButton > button:active { transform: translateY(0px) !important; }

/* ── Download buttons ── */
.stDownloadButton > button {
    background: var(--surface) !important;
    border: 1px solid var(--border-lit) !important;
    color: var(--text) !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.06em !important;
    height: 2.6rem !important;
    border-radius: 10px !important;
    width: 100% !important;
    transition: all 0.2s !important;
}
.stDownloadButton > button:hover {
    background: var(--surface-hover) !important;
    border-color: var(--orange) !important;
    color: var(--orange) !important;
}

/* ── Pipeline sidebar cards ── */
.pipe-card {
    border: 1px solid var(--border);
    border-radius: 14px;
    padding: 1.1rem 1.3rem;
    margin-bottom: 0.75rem;
    background: var(--surface);
    transition: all 0.4s cubic-bezier(0.4,0,0.2,1);
    position: relative;
    overflow: hidden;
}
.pipe-card::before {
    content: '';
    position: absolute;
    left: 0; top: 0; bottom: 0;
    width: 3px;
    background: var(--border);
    transition: background 0.4s;
}
.pipe-card.active::before { background: var(--orange); }
.pipe-card.done::before   { background: var(--green); }

.pipe-card.active {
    border-color: rgba(255,140,50,0.3);
    background: var(--orange-dim);
    box-shadow: 0 0 30px rgba(255,140,50,0.1);
}
.pipe-card.done {
    border-color: rgba(74,222,128,0.2);
    background: var(--green-dim);
}

.pipe-num {
    font-family: 'DM Mono', monospace;
    font-size: 0.6rem;
    color: var(--muted);
    margin-bottom: 4px;
}
.pipe-name {
    font-family: 'Syne', sans-serif;
    font-weight: 700;
    font-size: 0.92rem;
    color: var(--text);
    margin-bottom: 2px;
}
.pipe-sub {
    font-size: 0.72rem;
    color: var(--muted);
}
.pipe-indicator {
    position: absolute;
    right: 1.1rem;
    top: 50%;
    transform: translateY(-50%);
    width: 8px;
    height: 8px;
    border-radius: 50%;
    background: var(--muted2);
}
.pipe-indicator.running {
    background: var(--orange);
    box-shadow: 0 0 10px var(--orange);
    animation: blink 1.4s infinite;
}
.pipe-indicator.done {
    background: var(--green);
}

@keyframes blink {
    0%,100% { opacity: 1; }
    50% { opacity: 0.3; }
}

/* ── Report output ── */
.report-container {
    background: var(--surface);
    border: 1px solid var(--border-lit);
    border-radius: 20px;
    padding: 2.5rem 3rem;
    margin-top: 2rem;
}
.report-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 1.5rem;
    padding-bottom: 1rem;
    border-bottom: 1px solid var(--border);
}
.report-title-tag {
    font-family: 'DM Mono', monospace;
    font-size: 0.62rem;
    letter-spacing: 0.2em;
    color: var(--orange);
    text-transform: uppercase;
}
.report-timestamp {
    font-family: 'DM Mono', monospace;
    font-size: 0.6rem;
    color: var(--muted);
}

/* ── Critique expander ── */
.streamlit-expanderHeader {
    font-family: 'DM Mono', monospace !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.12em !important;
    color: var(--muted) !important;
    background: transparent !important;
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
}

/* ── Section divider ── */
.divider {
    height: 1px;
    background: var(--border);
    margin: 2rem 0;
}

/* ── Stats strip ── */
.stats-strip {
    display: flex;
    gap: 2rem;
    margin-bottom: 2.5rem;
}
.stat-item { text-align: left; }
.stat-num {
    font-family: 'Syne', sans-serif;
    font-weight: 800;
    font-size: 1.6rem;
    color: var(--orange);
    line-height: 1;
}
.stat-label {
    font-family: 'DM Mono', monospace;
    font-size: 0.58rem;
    letter-spacing: 0.15em;
    color: var(--muted);
    text-transform: uppercase;
    margin-top: 2px;
}

/* ── Pipeline section header ── */
.section-label {
    font-family: 'DM Mono', monospace;
    font-size: 0.6rem;
    letter-spacing: 0.25em;
    text-transform: uppercase;
    color: var(--muted);
    margin-bottom: 1rem;
    padding-bottom: 0.5rem;
    border-bottom: 1px solid var(--border);
}

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: var(--muted2); border-radius: 3px; }
</style>
""", unsafe_allow_html=True)

# ── Top Bar ───────────────────────────────────────────────────────────────────
st.markdown("""
<div class="topbar">
    <div class="topbar-logo">Research<span>Mind</span></div>
    <div class="topbar-badge">Multi-Agent Engine v2.0</div>
</div>
""", unsafe_allow_html=True)

# ── Main Layout ───────────────────────────────────────────────────────────────
left, right = st.columns([7, 4], gap="large")

with left:
    # Hero
    st.markdown('<div class="hero-eyebrow">Intelligence Orchestration Platform</div>', unsafe_allow_html=True)
    st.markdown('<div class="hero-title">Deep Research.<br><span class="accent">On Demand.</span></div>', unsafe_allow_html=True)
    st.markdown('<div class="hero-sub">A four-stage multi-agent pipeline that searches, reads, synthesises, and peer-reviews any topic into a structured technical report — with one click.</div>', unsafe_allow_html=True)

    # Stats
    st.markdown("""
    <div class="stats-strip">
        <div class="stat-item"><div class="stat-num">4</div><div class="stat-label">Agent Stages</div></div>
        <div class="stat-item"><div class="stat-num">∞</div><div class="stat-label">Research Depth</div></div>
        <div class="stat-item"><div class="stat-num">2</div><div class="stat-label">Export Formats</div></div>
    </div>
    """, unsafe_allow_html=True)

    # Input
    st.markdown('<div class="input-wrapper">', unsafe_allow_html=True)
    st.markdown('<div class="input-label">Research Objective</div>', unsafe_allow_html=True)
    topic = st.text_input(
        label="hidden",
        label_visibility="collapsed",
        placeholder="e.g. Impact of solid-state batteries on EV range by 2030...",
        key="topic_input"
    )
    st.markdown('</div>', unsafe_allow_html=True)

    if st.button("Initialize Research Sequence →", use_container_width=True):
        if topic.strip():
            st.session_state.results = {}
            st.session_state.stage = "searching"
            st.session_state.topic_submitted = topic.strip()
            st.rerun()
        else:
            st.error("Please enter a research objective.")

    # ── Results ──
    if st.session_state.stage == "complete" or "writer" in st.session_state.results:
        report_md   = st.session_state.results.get("writer", "")
        critique_md = st.session_state.results.get("critic", "")
        topic_used  = st.session_state.topic_submitted
        ts_display  = datetime.now().strftime("%d %b %Y, %H:%M")

        st.markdown(f"""
        <div class="report-container">
            <div class="report-header">
                <div class="report-title-tag">📄 Synthesis Report</div>
                <div class="report-timestamp">{ts_display} UTC</div>
            </div>
        """, unsafe_allow_html=True)

        st.markdown(report_md)

        if critique_md:
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            with st.expander("🧐 PEER REVIEW & CRITIQUE"):
                st.markdown(critique_md)

        st.markdown('</div>', unsafe_allow_html=True)

        # ── Export ──
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Export Report</div>', unsafe_allow_html=True)
        ec1, ec2 = st.columns(2)

        with ec1:
            try:
                pdf_bytes = generate_pdf(topic_used, report_md, critique_md)
                st.download_button(
                    label="⬇ Download PDF",
                    data=pdf_bytes,
                    file_name=f"researchmind_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"PDF generation failed: {e}")

        with ec2:
            try:
                docx_bytes = generate_docx(topic_used, report_md, critique_md)
                st.download_button(
                    label="⬇ Download DOCX",
                    data=docx_bytes,
                    file_name=f"researchmind_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"DOCX generation failed: {e}")


# ── Right Column: Pipeline Tracker ───────────────────────────────────────────
with right:
    st.markdown('<div class="section-label">Pipeline Status</div>', unsafe_allow_html=True)

    STAGES = [
        ("searching", "01", "Search Agent",   "Web indexing & source discovery"),
        ("reading",   "02", "Reader Agent",   "Deep content extraction"),
        ("writing",   "03", "Writer Chain",   "Narrative synthesis"),
        ("critiquing","04", "Critic Chain",   "Peer review & quality assurance"),
    ]

    DONE_KEYS = {
        "searching":  "search",
        "reading":    "reader",
        "writing":    "writer",
        "critiquing": "critic",
    }

    for stage_key, num, title, sub in STAGES:
        done    = DONE_KEYS[stage_key] in st.session_state.results
        active  = st.session_state.stage == stage_key
        card_cls  = "active" if active else ("done" if done else "")
        ind_cls   = "running" if active else ("done" if done else "")
        icon      = "✓" if done else ("▶" if active else "·")

        st.markdown(f"""
        <div class="pipe-card {card_cls}">
            <div class="pipe-num">{num} — {'COMPLETE' if done else ('RUNNING' if active else 'QUEUED')}</div>
            <div class="pipe-name">{title}</div>
            <div class="pipe-sub">{sub}</div>
            <div class="pipe-indicator {ind_cls}"></div>
        </div>
        """, unsafe_allow_html=True)

    # Progress summary
    done_count = sum(1 for k in DONE_KEYS.values() if k in st.session_state.results)
    if done_count > 0:
        st.markdown(f"""
        <div style="margin-top:1.5rem; padding:1rem 1.2rem; border:1px solid var(--border);
                    border-radius:12px; background: var(--surface);">
            <div style="font-family:'DM Mono'; font-size:0.6rem; color:var(--muted); margin-bottom:0.5rem;">
                COMPLETION
            </div>
            <div style="font-family:'Syne'; font-weight:800; font-size:1.8rem; color:var(--orange);">
                {done_count}/4
            </div>
            <div style="margin-top:0.5rem; height:4px; background:var(--border); border-radius:2px;">
                <div style="height:4px; width:{done_count*25}%; background:var(--orange);
                            border-radius:2px; transition:width 0.5s;"></div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    # Info block
    st.markdown("""
    <div style="margin-top:1.5rem; padding:1.2rem; border:1px solid var(--border);
                border-radius:12px; background:var(--surface);">
        <div style="font-family:'DM Mono'; font-size:0.6rem; color:var(--orange); margin-bottom:0.8rem; letter-spacing:0.15em;">
            HOW IT WORKS
        </div>
        <div style="font-size:0.8rem; color:var(--muted); line-height:1.7;">
            <b style="color:var(--text)">Search Agent</b> discovers authoritative sources across the web.<br><br>
            <b style="color:var(--text)">Reader Agent</b> extracts and structures raw content from each source.<br><br>
            <b style="color:var(--text)">Writer Chain</b> synthesises findings into a structured technical report.<br><br>
            <b style="color:var(--text)">Critic Chain</b> performs automated peer review for accuracy and completeness.
        </div>
    </div>
    """, unsafe_allow_html=True)


# ── Agent Execution ───────────────────────────────────────────────────────────
if st.session_state.stage not in ("idle", "complete"):
    topic_run = st.session_state.topic_submitted or topic

    try:
        from agents import build_reader_agent, build_search_agent, writer_chain, critic_chain

        if st.session_state.stage == "searching":
            agent = build_search_agent()
            res = agent.invoke({"messages": [("user", f"Find deep technical info on: {topic_run}")]})
            st.session_state.results["search"] = res["messages"][-1].content
            st.session_state.stage = "reading"
            st.rerun()

        elif st.session_state.stage == "reading":
            agent = build_reader_agent()
            res = agent.invoke({"messages": [("user", f"Context: {st.session_state.results['search'][:1200]}")]})
            st.session_state.results["reader"] = res["messages"][-1].content
            st.session_state.stage = "writing"
            st.rerun()

        elif st.session_state.stage == "writing":
            res = writer_chain.invoke({
                "topic": topic_run,
                "research": f"{st.session_state.results['search']}\n\n{st.session_state.results['reader']}"
            })
            st.session_state.results["writer"] = res
            st.session_state.stage = "critiquing"
            st.rerun()

        elif st.session_state.stage == "critiquing":
            res = critic_chain.invoke({"report": st.session_state.results["writer"]})
            st.session_state.results["critic"] = res
            st.session_state.stage = "complete"
            st.rerun()

    except ImportError:
        # Demo mode: agents.py not present — show placeholder output
        import random, time
        time.sleep(0.6)

        placeholder_data = {
            "searching": ("search", "reading",
                "**Search results:** Found 12 authoritative sources on the topic including IEEE papers, industry whitepapers, and recent news articles."),
            "reading": ("reader", "writing",
                "**Extracted content:** Synthesised key findings from 8 high-quality sources covering technical specifications, market analysis, and expert projections."),
            "writing": ("writer", "critiquing",
                f"# {topic_run}\n\n## Executive Summary\nThis report synthesises current intelligence on **{topic_run}**, drawing from 12 web sources and 8 deeply read documents.\n\n## Key Findings\n- **Finding 1:** Industry adoption is accelerating, with major players committing significant R&D budgets.\n- **Finding 2:** Technical barriers remain but are projected to resolve within 3–5 years.\n- **Finding 3:** Market projections suggest a 40–60% improvement in core metrics by 2030.\n\n## Technical Analysis\nCurrent state-of-the-art demonstrates promising trajectories. Peer-reviewed literature supports the optimistic mid-term outlook, while acknowledging manufacturing scale challenges.\n\n## Market Outlook\nEarly-adopter organisations are reporting measurable gains. The competitive landscape is shifting toward vertically integrated solutions.\n\n## Conclusion\nThe evidence supports a cautiously optimistic outlook. Organisations should begin strategic planning for integration within the next 18–24 months."),
            "critiquing": ("critic", "complete",
                "## Peer Review Summary\n\n**Overall Assessment:** The report is well-structured and grounded in available evidence.\n\n### Strengths\n- Clear executive summary with actionable conclusions\n- Balanced treatment of both optimistic and cautious perspectives\n- Strong citation of technical literature\n\n### Areas for Improvement\n- Regional market differences could be explored in more depth\n- Quantitative projections should include confidence intervals\n- A dedicated risk section would strengthen the analysis\n\n**Verdict:** Approved for distribution with minor revisions recommended."),
        }

        if st.session_state.stage in placeholder_data:
            key, next_stage, content = placeholder_data[st.session_state.stage]
            st.session_state.results[key] = content
            st.session_state.stage = next_stage
            st.rerun()


# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<div style="margin-top:5rem; padding-top:1.5rem; border-top:1px solid var(--border);
            display:flex; justify-content:space-between; align-items:center;">
    <div style="font-family:'DM Mono'; font-size:0.58rem; color:var(--muted2); letter-spacing:0.1em;">
        RESEARCHMIND ENGINE v2.0 &nbsp;·&nbsp; MULTI-AGENT ARCHITECTURE
    </div>
    <div style="font-family:'DM Mono'; font-size:0.58rem; color:var(--muted2);">
        SEARCH → READ → WRITE → REVIEW
    </div>
</div>
""", unsafe_allow_html=True)